import docx, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
doc = docx.Document(r'temp_zip_content\Victorem_SDD_v2.1.docx')

print('=== TABLES ===')
for ti, table in enumerate(doc.tables):
    print('')
    print('Table ' + str(ti+1) + ':')
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print('  Row ' + str(ri) + ': ' + ' | '.join(cells))

print('')
print('=== SECTIONS 5-8 FULL TEXT ===')
capture = False
for p in doc.paragraphs:
    txt = p.text.strip()
    if '5.' in txt and 'ARQUITECTURA' in txt:
        capture = True
    if 'VICTOREM' in txt and 'SDD' in txt and 'v2.1' in txt:
        break
    if capture and txt:
        print(txt)
